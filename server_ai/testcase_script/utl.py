import os
import json
import ast
from docx import Document
from docx.shared import Pt
import io
from io import BytesIO
from django.http import HttpResponse
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def generate_and_download_docx(request, generated_testcases, inputdocument_name):

    

        #print(generated_testcases)
        
        doc = Document()
        doc.add_heading('Testcases', 0)

        # set default font
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(10.5)
        font.name = 'Arial'

        # create table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # table headers
        headers = ["No.", "Testpoint", "Operation", "Expected Result"]
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            para = hdr_cells[i].paragraphs[0]
            run = para.add_run(text)
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # iterate over generated testcases and fill the table
        case_no = 1
        for item in generated_testcases:
            if isinstance(item, list) and len(item) > 0:
                raw = item[0]
            elif isinstance(item, str):
                raw = item
            else:
                continue

            
            def extract(pattern):
                m = re.search(pattern, raw)
                return m.group(1).replace("\\n", "\n") if m else ""

            testpoint = extract(r"'testpoint':\s*'([^']*)'")
            operation = extract(r"'operation':\s*'([^']*)'")
            expected = extract(r"'expectedresult':\s*'([^']*)'")

            
            row_cells = table.add_row().cells
            row_cells[0].text = str(case_no)  # No.
            row_cells[1].text = testpoint
            row_cells[2].text = operation
            row_cells[3].text = expected

            # set cell style
            for cell in row_cells:
                for para in cell.paragraphs:
                    para.style = doc.styles['Normal']

            case_no += 1

        # set column widths
        table.columns[0].width = Inches(0.6)  # No.
        table.columns[1].width = Inches(2)    # Testpoint
        table.columns[2].width = Inches(3.5)  # Operation
        table.columns[3].width = Inches(2.5)  # Expected Result

        # output the document to a BytesIO stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = 'attachment; filename=\"testcases.docx\"'
        return response

def extract_outer_braces(text):
    """
    提取最外层花括号 {} 中的内容
    """
    pattern = r'\{[^{}]*\}'  # 匹配不包含其他花括号的最外层花括号
    match = re.search(pattern, text)
    if match:
        return match.group(0)
    return None

def remove_json_comments(json_str):
    """
    移除 JSON 字符串中的 JavaScript 注释
    """
    # 移除单行注释
    json_str = re.sub(r'//.*', '', json_str)
    # 移除多行注释
    json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)
    return json_str

def clean_json_text(text):
    
    cleaned_lines = []
    for line in text.splitlines():
        # 去掉前后空格后判断是不是注释行
        stripped = line.strip()
        if stripped.startswith("#"):
            continue  # 整行注释，跳过
        # 如果一行里面带 #，只保留前面的部分
        if "#" in line:
            line = line.split("#", 1)[0].rstrip()
        if line.strip():  # 避免空行
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

def clean_and_parse_json(json_str):
    """
    清理和解析 JSON 字符串，处理常见的格式问题
    """
    # 1. 移除 JavaScript 注释
    json_str = re.sub(r'//.*', '', json_str)
    json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)
    
    # 2. 修复字符串中的引号问题
    def escape_quotes_in_string(match):
        content = match.group(1)
        # 转义内容中的双引号，但保留已经转义的引号
        content = re.sub(r'(?<!\\)"', r'\\"', content)
        return f'"{content}"'
    
    # 3. 应用引号转义到所有字符串值
    json_str = re.sub(r'"([^"]*)"', escape_quotes_in_string, json_str)
    
    # 4. 修复其他常见问题
    # 移除尾随逗号
    json_str = re.sub(r',\s*}', '}', json_str)
    json_str = re.sub(r',\s*]', ']', json_str)
    
    # 5. 尝试解析
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"JSON 解析错误: {e}")
        print(f"错误位置: 第{e.lineno}行第{e.colno}列")
        
        # 显示错误行
        lines = json_str.split('\n')
        if e.lineno <= len(lines):
            error_line = lines[e.lineno - 1]
            print(f"错误行: {error_line}")
            print(f"错误位置: {' ' * (e.colno - 1)}^")
        
        raise

def get_docx_files(directory_path):
    """使用列表推导式获取docx文件"""
    if not os.path.exists(directory_path):
        return []
    
    return [f for f in os.listdir(directory_path) 
            if f.lower().endswith('.docx')]
    
