import jieba
from docx import Document   
import re

class WordDocumentParser:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.document = Document(doc_path)
        self.paragraphs_content = [] 
        # 加载Jieba
        #jieba.enable_paddle()  # 启用Paddle模式，可以基于Paddle模型进行分词

    def parse(self):
        sentences = []
        for para in self.document.paragraphs:
            if 'List Paragraph' in para.style.name or 'Normal' in para.style.name:
                sentences.extend(self.split_sentences(para.text))
        return sentences

    def split_sentences(self, text):
        # split text into sentences using jieba
        words = jieba.cut(text, use_paddle=True)
        # convert the result to a string
        text = ''.join(words)
        text_data = re.split(r'(?<=[.!?。！？])', text)
        test_data_list = []
        for data in text_data:
            if data == '':
                continue
            if len(data) < 4:
                continue
            test_data_list.append(data)
        return test_data_list
    
    def get_paragraphs(self):
        # get all paragraphs with their headings
        Heading = ''
        
        for para in self.document.paragraphs:
            if 'Heading' in para.style.name:
                if len(para.text) != 0:
                    Heading = para.text

            if Heading != para.text and para.text != '':
                self.paragraphs_content.append({Heading: para.text})

        return self.paragraphs_content



